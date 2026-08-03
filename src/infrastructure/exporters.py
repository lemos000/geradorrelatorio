import pandas as pd
import os
import sys
# from pptx import Presentation
# from pptx.util import Inches, Pt
# from pptx.enum.text import PP_ALIGN
# from pptx.dml.color import RGBColor
# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Inches as DocxInches, Pt as DocxPt
# pyrefly: ignore [missing-import]
from docx.enum.text import WD_ALIGN_PARAGRAPH
# pyrefly: ignore [missing-import]
from docx.oxml import OxmlElement
# pyrefly: ignore [missing-import]
from docx.oxml.ns import qn
# pyrefly: ignore [missing-import]
from docx.shared import RGBColor

class FileExporter:
    """Classe responsável por exportar dados para diversos formatos (Excel, PPT, Word)."""
    
    def _set_cell_background(self, cell, fill):
        """Define a cor de fundo de uma célula."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_border(self, cell):
        """Adiciona bordas pretas sólidas a uma célula e centraliza verticalmente."""
        tcPr = cell._tc.get_or_add_tcPr()
        
        # Bordas
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            edge_elm = OxmlElement(f'w:{edge}')
            edge_elm.set(qn('w:val'), 'single')
            edge_elm.set(qn('w:sz'), '4')  # 1/2 pt
            edge_elm.set(qn('w:space'), '0')
            edge_elm.set(qn('w:color'), '000000')
            tcBorders.append(edge_elm)
        tcPr.append(tcBorders)

        # Alinhamento vertical
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def _get_resource_path(self, relative_path):
        """ Retorna o caminho absoluto para recursos, funcionando em dev e no PyInstaller """
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    @staticmethod
    def format_value(val, col_name):
        if hasattr(val, 'strftime'):
            return val.strftime('%d/%m/%Y')
        if isinstance(val, str):
            try:
                dt = pd.to_datetime(val)
                if not pd.isna(dt):
                    return dt.strftime('%d/%m/%Y')
            except Exception:
                pass
        if 'ADS' in str(col_name):
            return f"{int(val)}"
        if any(k in str(col_name) for k in ['ADR', 'Receita', 'RevPAR']):
            return f"R$ {val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        if '%' in str(col_name) or 'Ocupação' in str(col_name):
            return f"{val:.2f}%"
        if 'Reservas' in str(col_name):
            return f"{int(val)}"
        if isinstance(val, (float, int)):
            return f"{val:.2f}"
        return str(val)

    def export_to_excel(self, dataframes, sheet_names, output_path):
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            for df, name in zip(dataframes, sheet_names):
                df.to_excel(writer, sheet_name=name, index=False)

    def _add_page_number(self, run):
        """Helper to add dynamic page numbering field."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def export_to_docx(self, dataframes, summaries, evol_df, closing_df, charts, output_path, property_name="[NOME DA PROPRIEDADE]", reference=""):
        doc = Document()
        
        # --- Global Style (Inter) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Inter'
        font.size = DocxPt(11)

        # --- Heading Styles (Poppins, Size 14, Bold, Centered) ---
        for level in range(1, 3):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Poppins'
            h_style.font.size = DocxPt(14)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 81, 81)  # Midnight Green #005151
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Footer Configuration ---
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        
        # Footer for Page 1 (Cover) - Only Copyright
        f_first = section.first_page_footer
        p_first = f_first.paragraphs[0]
        p_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f1 = p_first.add_run('Revinn Estratégias © Copyright – Todos os direitos reservados.')
        run_f1.font.name = 'Inter'
        run_f1.font.size = DocxPt(9)

        # Footer for Page 2 onwards - Copyright + Page Number
        f_rest = section.footer
        p_rest = f_rest.paragraphs[0]
        p_rest.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f2 = p_rest.add_run('Revinn Estratégias © Copyright – Todos os direitos reservados. | Página ')
        run_f2.font.name = 'Inter'
        run_f2.font.size = DocxPt(9)
        
        run_page = p_rest.add_run()
        run_page.font.name = 'Inter'
        run_page.font.size = DocxPt(9)
        self._add_page_number(run_page)

        # --- Cover ---
        # Logo at top
        logo_path = self._get_resource_path('image.png')
        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_path, width=DocxInches(1.5))

        for _ in range(2): doc.add_paragraph() # Reduzido de 4 para 2

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Revinn Estratégias')
        run.font.name = 'Poppins'
        run.font.size = DocxPt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Departamento de Consultoria')
        run.font.name = 'Inter'
        run.font.size = DocxPt(11)

        for _ in range(3): doc.add_paragraph() # Reduzido de 7 para 3

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Limpa o nome removendo underscores
        clean_name = property_name.replace('_', ' ').upper()
        run = p.add_run(clean_name)
        run.font.name = 'Poppins'
        run.font.size = DocxPt(18)
        run.font.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = 'RELATÓRIO DE KPI MENSAL'
        if reference:
            subtitle += f" - {reference}"
        run = p.add_run(subtitle)
        run.font.name = 'Poppins'
        run.font.size = DocxPt(14)

        # Push São Paulo / 2026 to bottom
        for _ in range(4): doc.add_paragraph() # Reduzido de 8 para 4

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('São Paulo\n2026')
        run.font.name = 'Inter'
        run.font.size = DocxPt(11)

        doc.add_page_break()

        # --- Audit Tables ---
        titles = ["Check-outs por Categoria", "Check-outs por Canal", "Reservas Emitidas por Canal"]
        # Relaciona cada tabela com seu respectivo resumo (sums[0] para CO, sums[1] para RE)
        table_sums = [summaries[0], summaries[0], summaries[1]]
        
        for i, df in enumerate(dataframes):
            summary_data = table_sums[i]
            heading = doc.add_heading(titles[i], level=1)
            
            self._add_docx_table(doc, df, summary=summary_data)
            
            # Adiciona descritivo abaixo da tabela
            revenue_str = self.format_value(summary_data.get('total_receita', 0), 'Receita')
            nights = summary_data.get('total_noites', 0)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"RN's: {nights}; Renda Total {revenue_str}")
            run.font.name = 'Inter'
            run.font.bold = True
            run.font.size = DocxPt(9)
            
            doc.add_page_break()

        # --- Climber | YoY Mensal (Análise de Pickup) ---
        if evol_df is not None and not evol_df.empty:
            heading = doc.add_heading('Climber | YoY Mensal', level=1)
            
            if len(charts) >= 3:
                heading = doc.add_heading('Evolução Diária (Linhas)', level=2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(charts[0], width=DocxInches(6.5))
                
                heading = doc.add_heading('Comparativo de Fechamento', level=2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(charts[1], width=DocxInches(3.1))
                p.add_run().add_picture(charts[2], width=DocxInches(3.1))
                
                doc.add_page_break()
                self._add_closing_docx_table(doc, closing_df)
                doc.add_page_break()
                self._add_evolution_docx_table(doc, evol_df)
            else:
                doc.add_paragraph("Dados de evolução insuficientes para gerar gráficos.")
                if closing_df is not None and not closing_df.empty:
                    self._add_closing_docx_table(doc, closing_df)
        else:
            # Caso não haja dados de pickup, apenas adiciona a página vazia padrão
            doc.add_page_break()
            heading = doc.add_heading('Climber | YoY Mensal', level=1)

        # --- Additional Empty Pages ---
        extra_pages = [
            "Notas & Ranking (Last 90) | Booking.com",
            "Performance (Last 90) | Expedia"
        ]
        for title in extra_pages:
            doc.add_page_break()
            heading = doc.add_heading(title, level=1)

        # --- Back Cover ---
        doc.add_page_break()
        for _ in range(10): doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Revinn Estratégias')
        run.font.name = 'Poppins'
        run.font.size = DocxPt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Consultoria e Estratégia Hoteleira')
        run.font.name = 'Inter'
        run.font.size = DocxPt(11)
        
        doc.save(output_path)

    def set_table_autofit(self, table):
        """Ajusta a tabela para se auto-ajustar ao conteúdo (AutoFit to Contents)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = tblPr.xpath('w:tblW')
        if tblW:
            tblW[0].set(qn('w:type'), 'auto')
            tblW[0].set(qn('w:w'), '0')

    def _add_docx_table(self, doc, df, summary=None):
        rows_count = len(df) + 1
        if summary: rows_count += 1
        
        table = doc.add_table(rows=rows_count, cols=len(df.columns))
        self.set_table_autofit(table)
        
        # Header (Cor: #159F92, Texto: Branco)
        for i, col in enumerate(df.columns): 
            cell = table.rows[0].cells[i]
            cell.text = col
            self._set_cell_background(cell, "159F92")
            self._set_cell_border(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Poppins'
                    run.font.size = DocxPt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data Rows
        for i, row in enumerate(df.values):
            bg_color = "E8E8E8" if i % 2 == 0 else "FFFFFF"
            for j, val in enumerate(row):
                cell = table.rows[i+1].cells[j]
                cell.text = self.format_value(val, df.columns[j])
                self._set_cell_border(cell)
                
                if j == 0:
                    self._set_cell_background(cell, "159F92")
                else:
                    self._set_cell_background(cell, bg_color)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Inter'
                        run.font.size = DocxPt(9)
                        if j == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Summary Row (TOTAL)
        if summary:
            last_row = table.rows[-1].cells
            last_row[0].text = "TOTAL"
            self._set_cell_background(last_row[0], "159F92")
            
            for j, cell_item in enumerate(last_row):
                self._set_cell_border(cell_item)
                if j > 0:
                    self._set_cell_background(cell_item, "E8E8E8" if len(df) % 2 == 0 else "FFFFFF")
                
                if j == 0: 
                    for p in last_row[0].paragraphs:
                        for run in p.runs: 
                            run.font.name = 'Inter'
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    continue
                
                col_name = df.columns[j]
                
                if 'Receita' in col_name or 'ADR' in col_name or 'ADS' in col_name:
                    if 'Receita' in col_name: 
                        val = summary.get('total_receita', 0)
                        last_row[j].text = self.format_value(val, col_name)
                    elif 'ADR' in col_name:
                        rev = summary.get('total_receita', 0)
                        nights = summary.get('total_noites', 1)
                        val = rev / nights if nights > 0 else 0
                        last_row[j].text = self.format_value(val, col_name)
                    elif 'ADS' in col_name:
                        val = df[col_name].mean()
                        last_row[j].text = f"{val:.1f}".replace('.', ',')
                elif 'Reservas' in col_name:
                    val = summary.get('total_reservas', 0)
                    last_row[j].text = str(val)
                elif 'LOS' in col_name:
                    val = summary.get('total_noites', 0) / summary.get('total_reservas', 1)
                    last_row[j].text = f"{val:.2f}"
                elif 'Share' in col_name:
                    last_row[j].text = "100,00%"

            for j, cell in enumerate(last_row):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Inter'
                        run.font.bold = True
                        run.font.size = DocxPt(9)
                        if j == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)

    def _add_closing_docx_table(self, doc, df):
        heading = doc.add_heading("Tabela de Fechamento Consolidado", level=2)
        
        metrics = [
            ('Ocupação (%)', 'Ocupação Presente', 'Ocupação Passado', '{:.2f}%'),
            ('ADR (R$)', 'ADR Presente', 'ADR Passado', 'R$ {:,.2f}'),
            ('Receita (R$)', 'Receita Presente', 'Receita Passado', 'R$ {:,.2f}')
        ]
        table = doc.add_table(rows=1, cols=4)
        self.set_table_autofit(table)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Métrica', 'Ano Atual', 'Ano Passado', 'Variação (%)'
        for cell in hdr:
            self._set_cell_background(cell, "159F92")
            self._set_cell_border(cell)
            for p in cell.paragraphs:
                for run in p.runs: 
                    run.font.name = 'Poppins'
                    run.font.size = DocxPt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        for i, (label, pres, past, fmt) in enumerate(metrics):
            row = table.add_row().cells
            p_val, pas_val = df[pres].iloc[0], df[past].iloc[0]
            var = ((p_val - pas_val) / pas_val * 100) if pas_val != 0 else 0
            row[0].text = label
            row[1].text = fmt.format(p_val).replace(',', 'X').replace('.', ',').replace('X', '.')
            row[2].text = fmt.format(pas_val).replace(',', 'X').replace('.', ',').replace('X', '.')
            row[3].text = f"{var:+.2f}%"
            
            bg_color = "E8E8E8" if i % 2 == 0 else "FFFFFF"
            for j, cell in enumerate(row):
                self._set_cell_border(cell)
                if j == 0:
                    self._set_cell_background(cell, "159F92")
                else:
                    self._set_cell_background(cell, bg_color)
                
                for p in cell.paragraphs:
                    for run in p.runs: 
                        run.font.name = 'Inter'
                        run.font.size = DocxPt(9)
                        if j == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("* Nota: A Tabela de Fechamento Consolidado baseia-se na comparação do mesmo período (mês/dias equivalentes) entre o ano atual e o ano anterior (Year-over-Year).")
        run.font.name = 'Inter'
        run.font.italic = True
        run.font.size = DocxPt(8)

    def _add_evolution_docx_table(self, doc, df):
        heading = doc.add_heading("Detalhamento Diário YoY", level=2)
        
        cols = ['Data', 'Tipo', 'Ocupação (%)', 'ADR (R$)', 'RevPAR (R$)']
        table = doc.add_table(rows=1, cols=len(cols))
        self.set_table_autofit(table)
        for i, col in enumerate(cols): 
            cell = table.rows[0].cells[i]
            cell.text = col
            self._set_cell_background(cell, "159F92")
            self._set_cell_border(cell)
            for p in cell.paragraphs:
                for run in p.runs: 
                    run.font.name = 'Poppins'
                    run.font.size = DocxPt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        for i, (_, row) in enumerate(df.iterrows()):
            r1 = table.add_row().cells
            r1[0].text, r1[1].text, r1[2].text = self.format_value(row['Data'], 'Data'), "Ano Atual", f"{row['Ocupação Presente']:.2f}%"
            r1[3].text = self.format_value(row['ADR Presente'], 'ADR')
            r1[4].text = self.format_value(row['RevPAR Presente'], 'RevPAR')
            
            r2 = table.add_row().cells
            r2[0].text, r2[1].text, r2[2].text = self.format_value(row['Data'], 'Data'), "Ano Passado", f"{row['Ocupação Passado']:.2f}%"
            r2[3].text = self.format_value(row['ADR Passado'], 'ADR')
            r2[4].text = self.format_value(row['RevPAR Passado'], 'RevPAR')

            bg_color_pair = "E8E8E8" if i % 2 == 0 else "FFFFFF"
            
            for row_cells in [r1, r2]:
                for j, cell in enumerate(row_cells):
                    self._set_cell_border(cell)
                    if j == 0:
                        self._set_cell_background(cell, "159F92")
                    else:
                        self._set_cell_background(cell, bg_color_pair)
                    
                    for p in cell.paragraphs:
                        for run in p.runs: 
                            run.font.name = 'Inter'
                            run.font.size = DocxPt(9)
                            if j == 0:
                                run.font.color.rgb = RGBColor(255, 255, 255)

    def format_decimal(self, val):
        if pd.isna(val):
            return "0,00"
        try:
            return f"{float(val):.2f}".replace('.', ',')
        except (ValueError, TypeError):
            return str(val)

    def export_to_pdf(self, dataframes, summaries, evol_df, closing_df, output_path, property_name="[NOME DA PROPRIEDADE]", reference="", additional_sections=None):
        # Resolve logo.png URL
        logo_path = os.path.abspath("logo.png")
        logo_url = "file:///" + logo_path.replace("\\", "/") if os.path.exists(logo_path) else ""

        # 1. Configurar caminhos do DLL do GTK3 no Windows para evitar falhas do WeasyPrint
        if os.name == 'nt':
            possible_gtk_paths = [
                r"C:\Program Files\GTK3-Runtime Win64\bin",
                r"C:\Program Files (x86)\GTK3-Runtime Win64\bin",
                r"C:\msys64\mingw64\bin",
                r"C:\tools\msys64\mingw64\bin"
            ]
            for path in possible_gtk_paths:
                if os.path.exists(path):
                    try:
                        os.add_dll_directory(path)
                    except AttributeError:
                        os.environ['PATH'] = path + os.path.pathsep + os.environ['PATH']
                    break

        # Importa weasyprint localmente para evitar erros de importação na inicialização da aplicação
        try:
            from weasyprint import HTML
        except OSError as e:
            raise ImportError(
                "Não foi possível inicializar o WeasyPrint. Certifique-se de que o GTK+3 Runtime está instalado no Windows.\n"
                "Para resolver, acesse: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows\n"
                f"Erro detalhado: {str(e)}"
            )

        # 2. Obter valores para os cartões de KPI
        receita_val = "R$ 0,00"
        receita_var = "0,00%"
        receita_var_class = "positive"
        
        ocupacao_val = "0,00%"
        ocupacao_var = "0,00%"
        ocupacao_var_class = "positive"
        
        adr_val = "R$ 0,00"
        adr_var = "0,00%"
        adr_var_class = "positive"
        
        revpar_val = "R$ 0,00"
        # Função interna para montar o bloco de KPI no PDF
        def make_kpi_card(label, pres_val_str, var_str, var_class, past_val_str=None):
            past_html = ""
            if past_val_str is not None:
                past_html = f"""
                <hr class="kpi-separator" />
                <span class="kpi-past-label">Ano Anterior</span>
                <p class="kpi-past-value">{past_val_str}</p>
                """
            return f"""
            <div class="kpi-card">
                <span class="kpi-label">{label}</span>
                <p class="kpi-value">{pres_val_str}</p>
                <span class="kpi-sub {var_class}">{var_str}</span>
                {past_html}
            </div>
            """

        if closing_df is not None and not closing_df.empty:
            # Dados consolidados de fechamento do Pickup
            pres_rec = closing_df['Receita Presente'].iloc[0]
            past_rec = closing_df['Receita Passado'].iloc[0]
            receita_val = self.format_value(pres_rec, 'Receita')
            if past_rec != 0:
                var = (pres_rec - past_rec) / past_rec * 100
                receita_var = f"{var:+.2f}%".replace('.', ',')
                receita_var_class = "positive" if var >= 0 else "negative"
            else:
                receita_var = "0,00%"
                receita_var_class = "positive"
            past_rec_str = self.format_value(past_rec, 'Receita') if (past_rec is not None and past_rec > 0) else None

            pres_occ = closing_df['Ocupação Presente'].iloc[0]
            past_occ = closing_df['Ocupação Passado'].iloc[0]
            ocupacao_val = f"{pres_occ:.2f}%".replace('.', ',')
            # Variação em pontos percentuais (p.p.)
            var_pp = pres_occ - past_occ
            ocupacao_var = f"{var_pp:+.2f}".replace('.', ',') + " p.p."
            ocupacao_var_class = "positive" if var_pp >= 0 else "negative"
            past_occ_str = f"{past_occ:.2f}%".replace('.', ',') if (past_occ is not None and past_occ > 0) else None

            pres_adr = closing_df['ADR Presente'].iloc[0]
            past_adr = closing_df['ADR Passado'].iloc[0]
            adr_val = self.format_value(pres_adr, 'ADR')
            if past_adr != 0:
                var = (pres_adr - past_adr) / past_adr * 100
                adr_var = f"{var:+.2f}%".replace('.', ',')
                adr_var_class = "positive" if var >= 0 else "negative"
            else:
                adr_var = "0,00%"
                adr_var_class = "positive"
            past_adr_str = self.format_value(past_adr, 'ADR') if (past_adr is not None and past_adr > 0) else None

            # Se existirem colunas específicas de RevPAR, usá-las diretamente
            if 'RevPAR Presente' in closing_df.columns:
                pres_rev = closing_df['RevPAR Presente'].iloc[0]
                past_rev = closing_df['RevPAR Passado'].iloc[0]
            else:
                pres_rev = pres_occ * pres_adr / 100
                past_rev = past_occ * past_adr / 100
                
            revpar_val = self.format_value(pres_rev, 'RevPAR')
            if past_rev != 0:
                var = (pres_rev - past_rev) / past_rev * 100
                revpar_var = f"{var:+.2f}%".replace('.', ',')
                revpar_var_class = "positive" if var >= 0 else "negative"
            else:
                revpar_var = "0,00%"
                revpar_var_class = "positive"
            past_rev_str = self.format_value(past_rev, 'RevPAR') if (past_rev is not None and past_rev > 0) else None
        else:
            # Fallback para os dados de check-outs (CO)
            receita_val = "R$ 0,00"
            receita_var = "N/D"
            receita_var_class = "positive"
            past_rec_str = None
            
            ocupacao_val = "N/D"
            ocupacao_var = "N/D"
            ocupacao_var_class = "positive"
            past_occ_str = None
            
            adr_val = "R$ 0,00"
            adr_var = "N/D"
            adr_var_class = "positive"
            past_adr_str = None
            
            revpar_val = "N/D"
            revpar_var = "N/D"
            revpar_var_class = "positive"
            past_rev_str = None
            
            if summaries and len(summaries) > 0:
                co_sum = summaries[0]
                tot_rec = co_sum.get('total_receita', 0)
                tot_nights = co_sum.get('total_noites', 0)
                receita_val = self.format_value(tot_rec, 'Receita')
                ocupacao_var_class = "positive"
                
                calc_adr = tot_rec / tot_nights if tot_nights > 0 else 0
                adr_val = self.format_value(calc_adr, 'ADR')
                adr_var = "N/D"
                adr_var_class = "positive"
                
                revpar_val = "N/D"
                revpar_var = "N/D"
                revpar_var_class = "positive"

        # Monta os cartões de KPIs dinamicamente
        kpi_cards_html = f"""
        <div class="kpi-container">
            {make_kpi_card("Receita Líquida (YoY)", receita_val, receita_var, receita_var_class, past_rec_str)}
            {make_kpi_card("Taxa de Ocupação (YoY)", ocupacao_val, ocupacao_var, ocupacao_var_class, past_occ_str)}
            {make_kpi_card("Diária Média - ADR (YoY)", adr_val, adr_var, adr_var_class, past_adr_str)}
            {make_kpi_card("RevPAR Consolidado (YoY)", revpar_val, revpar_var, revpar_var_class, past_rev_str)}
        </div>
        """

        # 3. Construir linhas das tabelas em HTML
        # Tabela 1: Desempenho por Categoria de Acomodação (Check-outs)
        tabela_categorias_rows = []
        df_cat = dataframes[0]
        for _, row in df_cat.iterrows():
            tabela_categorias_rows.append(f"""
            <tr>
                <td>{row['Categoria']}</td>
                <td>{self.format_value(row['ADR'], 'ADR')}</td>
                <td>{self.format_decimal(row['LOS'])}</td>
                <td>{self.format_value(row['Receita'], 'Receita')}</td>
            </tr>
            """)
        if len(summaries) > 0:
            co_sum = summaries[0]
            tot_rec = co_sum.get('total_receita', 0)
            tot_nights = co_sum.get('total_noites', 0)
            tot_res = co_sum.get('total_reservas', 1)
            tot_adr = tot_rec / tot_nights if tot_nights > 0 else 0
            tot_los = tot_nights / tot_res if tot_res > 0 else 0
            tabela_categorias_rows.append(f"""
            <tr class="total-row">
                <td>TOTAL</td>
                <td>{self.format_value(tot_adr, 'ADR')}</td>
                <td>{self.format_decimal(tot_los)}</td>
                <td>{self.format_value(tot_rec, 'Receita')}</td>
            </tr>
            """)
        tabela_categorias_html = "\n".join(tabela_categorias_rows)

        # Tabela 2: Atribuição de Performance por Canais de Emissão (Check-outs)
        tabela_emissoes_passadas_rows = []
        df_canal_co = dataframes[1]
        for _, row in df_canal_co.iterrows():
            tabela_emissoes_passadas_rows.append(f"""
            <tr>
                <td>{row['Canal']}</td>
                <td>{self.format_value(row['Receita'], 'Receita')}</td>
                <td>{int(row['Reservas'])}</td>
                <td>{self.format_decimal(row['LOS'])}</td>
                <td>{self.format_value(row['ADR'], 'ADR')}</td>
                <td>{int(row['ADS'])}</td>
                <td>{self.format_decimal(row['Share (%)'])}%</td>
            </tr>
            """)
        if len(summaries) > 0:
            co_sum = summaries[0]
            tot_rec = co_sum.get('total_receita', 0)
            tot_res = co_sum.get('total_reservas', 0)
            tot_nights = co_sum.get('total_noites', 0)
            tot_los = tot_nights / tot_res if tot_res > 0 else 0
            tot_adr = tot_rec / tot_nights if tot_nights > 0 else 0
            tot_ads = df_canal_co['ADS'].mean() if not df_canal_co.empty else 0
            tabela_emissoes_passadas_rows.append(f"""
            <tr class="total-row">
                <td>TOTAL</td>
                <td>{self.format_value(tot_rec, 'Receita')}</td>
                <td>{tot_res}</td>
                <td>{self.format_decimal(tot_los)}</td>
                <td>{self.format_value(tot_adr, 'ADR')}</td>
                <td>{f"{tot_ads:.1f}".replace('.', ',')}</td>
                <td>100,00%</td>
            </tr>
            """)
        tabela_emissoes_passadas_html = "\n".join(tabela_emissoes_passadas_rows)

        # Tabela 3: Carteira de Captação Futura por Canais de Emissão (Pacing / OTB)
        tabela_emissoes_futuras_rows = []
        df_canal_re = dataframes[2]
        for _, row in df_canal_re.iterrows():
            tabela_emissoes_futuras_rows.append(f"""
            <tr>
                <td>{row['Canal']}</td>
                <td>{self.format_value(row['Receita'], 'Receita')}</td>
                <td>{int(row['Reservas'])}</td>
                <td>{self.format_value(row['ADR'], 'ADR')}</td>
                <td>{int(row['ADS'])}</td>
                <td>{self.format_decimal(row['Share (%)'])}%</td>
            </tr>
            """)
        if len(summaries) > 1:
            re_sum = summaries[1]
            tot_rec = re_sum.get('total_receita', 0)
            tot_res = re_sum.get('total_reservas', 0)
            tot_nights = re_sum.get('total_noites', 0)
            tot_adr = tot_rec / tot_nights if tot_nights > 0 else 0
            tot_ads = df_canal_re['ADS'].mean() if not df_canal_re.empty else 0
            tabela_emissoes_futuras_rows.append(f"""
            <tr class="total-row">
                <td>TOTAL</td>
                <td>{self.format_value(tot_rec, 'Receita')}</td>
                <td>{tot_res}</td>
                <td>{self.format_value(tot_adr, 'ADR')}</td>
                <td>{f"{tot_ads:.1f}".replace('.', ',')}</td>
                <td>100,00%</td>
            </tr>
            """)
        tabela_emissoes_futuras_html = "\n".join(tabela_emissoes_futuras_rows)

        # 3.5 Construir HTML para as seções adicionais dinâmicas
        additional_sections_html = ""
        if additional_sections:
            for sec in additional_sections:
                title = sec.get("title", "").strip()
                desc = sec.get("description", "").strip()
                img_path_raw = sec.get("image_path", "").strip()
                img_paths = [p.strip() for p in img_path_raw.split(";") if p.strip()]
                key = sec.get("key", "").lower().strip()
                
                if "booking" in key:
                    gallery_class = "booking-gallery"
                elif "expedia" in key:
                    gallery_class = "expedia-gallery"
                else:
                    gallery_class = "image-gallery"
                
                additional_sections_html += f"""
                <div style="page-break-before: always;">
                    <div class="table-row header-container">
                        <div class="table-cell" style="width: 50%; vertical-align: middle;">
                            <img src="{logo_url}" style="height: 100px; width: auto;" />
                        </div>
                        <div class="table-cell doc-info" style="width: 50%; vertical-align: middle;">
                            <h1>{property_name}</h1>
                            <p>Lâmina de Performance Mensal &bull; {reference}</p>
                        </div>
                    </div>
                """
                
                if title:
                    additional_sections_html += f"""
                    <div class="section-header">{title}</div>
                    """
                
                if desc:
                    formatted_desc = desc.replace("\n", "<br>")
                    additional_sections_html += f"""
                    <div style="font-size: 9.5pt; color: #2c3e50; line-height: 1.5; margin-bottom: 20px; text-align: justify; padding: 10px; background: #f8f9fa; border-left: 3px solid #159F92; border: 1px solid #e2e8f0; border-left-width: 3px; border-left-color: #159F92; border-radius: 0 4px 4px 0; font-family: 'Inter', sans-serif;">
                        {formatted_desc}
                    </div>
                    """
                
                if img_paths:
                    if gallery_class == "booking-gallery":
                        # Generate the simple but highly robust inline-block layout for Booking
                        resolved_urls = []
                        for img_p in img_paths:
                            if os.path.exists(img_p):
                                clean_path = os.path.abspath(img_p).replace("\\", "/")
                                if not clean_path.startswith("/"):
                                    clean_path = "/" + clean_path
                                resolved_urls.append(f"file://{clean_path}")
                        
                        if len(resolved_urls) >= 1:
                            img1 = resolved_urls[0]
                            img2 = resolved_urls[1] if len(resolved_urls) > 1 else ""
                            img3 = resolved_urls[2] if len(resolved_urls) > 2 else ""
                            
                            additional_sections_html += '\n<table class="booking-table-layout">'
                            additional_sections_html += f'\n    <tr><td colspan="2" style="padding-bottom: 8px !important;"><img src="{img1}" class="gallery-img-booking-1" /></td></tr>'
                            if img2 or img3:
                                additional_sections_html += '\n    <tr>'
                                if img2:
                                    additional_sections_html += f'<td style="width: 70% !important; padding-right: 4px !important;"><img src="{img2}" class="gallery-img-booking-2" /></td>'
                                if img3:
                                    additional_sections_html += f'<td style="width: 30% !important; padding-left: 4px !important;"><img src="{img3}" class="gallery-img-booking-3" /></td>'
                                additional_sections_html += '\n    </tr>'
                            additional_sections_html += '\n</table>'
                    else:
                        additional_sections_html += f'\n<div class="{gallery_class}">'
                        for idx, img_p in enumerate(img_paths):
                            if os.path.exists(img_p):
                                clean_path = os.path.abspath(img_p).replace("\\", "/")
                                if not clean_path.startswith("/"):
                                    clean_path = "/" + clean_path
                                file_url = f"file://{clean_path}"
                                
                                width_style = ""
                                if gallery_class == "image-gallery":
                                    try:
                                        # pyrefly: ignore [missing-import]
                                        from PIL import Image
                                        with Image.open(img_p) as im:
                                            w, h = im.size
                                            if w > 0:
                                                width_style = f"width: {w}px;"
                                    except Exception:
                                        pass
                                    
                                if gallery_class == "expedia-gallery":
                                    additional_sections_html += f'\n    <div class="grid-item-expedia"><img src="{file_url}" class="gallery-img" /></div>'
                                else:
                                    additional_sections_html += f'\n    <img src="{file_url}" class="gallery-img" style="{width_style}" />'
                        additional_sections_html += '\n</div>'
                
                additional_sections_html += """
                </div>
                """

        # 4. Definir HTML Template e substituir variáveis
        html_template = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <title>Revinn - Lâmina de Performance Mensal</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&family=Poppins:wght@400;600;700&display=swap');

        @page {{
            size: A4;
            margin: 15mm 12mm;
            background: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' width='800' height='1000' viewBox='0 0 800 1000'><text fill='rgba(240,240,240,0.45)' font-family='Inter, sans-serif' font-size='85' font-weight='bold' x='100' y='600' transform='rotate(-45 300 600)'>Revinn</text></svg>");
            @bottom-right {{
                content: "Pág. " counter(page);
                font-family: 'Inter', sans-serif;
                font-size: 8pt;
                color: #7f8c8d;
            }}
        }}

        @page :first {{
            @bottom-right {{
                content: "";
            }}
        }}
        
        *, *::before, *::after {{
            box-sizing: border-box;
        }}

        body {{
            font-family: 'Inter', sans-serif;
            color: #2c3e50;
            line-height: 1.4;
            margin: 0;
            padding: 0;
            background-color: #ffffff;
            font-size: 10pt;
        }}
        
        /* Cover Page Styling */
        .cover-page {{
            height: 92%;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            padding: 20mm 10mm;
            page-break-after: always;
            box-sizing: border-box;
        }}
        
        .cover-logo {{
            text-align: center;
            margin-bottom: 30px;
        }}
        
        .cover-brand-title {{
            font-family: 'Poppins', sans-serif;
            font-size: 32pt;
            font-weight: bold;
            color: #005151;
            text-transform: lowercase;
            display: block;
            letter-spacing: -1px;
        }}
        
        .cover-brand-subtitle {{
            font-family: 'Inter', sans-serif;
            font-size: 10pt;
            font-weight: bold;
            color: #7f8c8d;
            text-transform: uppercase;
            letter-spacing: 3px;
        }}
        
        .cover-main {{
            text-align: center;
            margin-top: 80px;
            margin-bottom: 80px;
        }}
        
        .cover-hotel-name {{
            font-family: 'Poppins', sans-serif;
            font-size: 26pt;
            color: #005151;
            text-transform: uppercase;
            font-weight: 700;
            margin: 0 0 10px 0;
            letter-spacing: 0.5px;
        }}
        
        .cover-doc-title {{
            font-family: 'Poppins', sans-serif;
            font-size: 14pt;
            color: #2c3e50;
            text-transform: uppercase;
            font-weight: 600;
            margin: 0 0 5px 0;
            letter-spacing: 1px;
        }}
        
        .cover-doc-reference {{
            font-family: 'Inter', sans-serif;
            font-size: 12pt;
            color: #7f8c8d;
            font-weight: 400;
            margin: 0;
        }}
        
        .cover-footer {{
            margin-top: auto;
            border-top: 1px solid #e2e8f0;
            padding-top: 20px;
            text-align: center;
        }}
        
        .cover-privacy-title {{
            font-family: 'Poppins', sans-serif;
            font-size: 8pt;
            font-weight: 700;
            color: #005151;
            margin-bottom: 8px;
            letter-spacing: 0.5px;
        }}
        
        .cover-privacy-text {{
            font-family: 'Inter', sans-serif;
            font-size: 7.5pt;
            color: #7f8c8d;
            line-height: 1.5;
            max-width: 500px;
            margin: 0 auto 20px auto;
            text-align: justify;
        }}
        
        .cover-date {{
            font-family: 'Inter', sans-serif;
            font-size: 9pt;
            color: #2c3e50;
            margin: 0;
        }}
        
        /* Estrutura de colunas via display: table para compatibilidade WeasyPrint */
        .table-row {{
            display: table;
            width: 100%;
            table-layout: fixed;
            margin-bottom: 15px;
        }}
        
        .table-cell {{
            display: table-cell;
            vertical-align: top;
        }}
        
        /* Cabeçalho Corporativo */
        .header-container {{
            border-bottom: 2px solid #159F92;
            padding-bottom: 8px;
            margin-bottom: 20px;
        }}
        
        .brand-title {{
            font-family: 'Poppins', sans-serif;
            font-size: 22pt;
            font-weight: bold;
            color: #005151;
            text-transform: lowercase;
            margin: 0;
            letter-spacing: -0.5px;
        }}
        
        .brand-subtitle {{
            font-family: 'Inter', sans-serif;
            font-size: 8pt;
            font-weight: bold;
            color: #7f8c8d;
            text-transform: uppercase;
            letter-spacing: 1.5px;
            margin-top: -2px;
        }}
        
        .doc-info {{
            text-align: right;
        }}
        
        .doc-info h1 {{
            font-family: 'Poppins', sans-serif;
            font-size: 14pt;
            margin: 0;
            color: #005151;
            text-transform: uppercase;
            font-weight: 700;
        }}
        
        .doc-info p {{
            font-family: 'Inter', sans-serif;
            font-size: 9pt;
            margin: 3px 0 0 0;
            color: #7f8c8d;
            font-weight: 600;
        }}

        /* Blocos de Indicadores Gerais (KPIs) */
        .kpi-container {{
            display: table;
            width: 100%;
            table-layout: fixed;
            margin-bottom: 20px;
            border-collapse: separate;
            border-spacing: 8px 0;
            margin-left: -8px;
            margin-right: -8px;
        }}
        
        .kpi-card {{
            display: table-cell;
            background: #ffffff;
            border: 1px solid #e2e8f0;
            border-top: 3px solid #159F92;
            padding: 12px 10px;
            border-radius: 4px;
            text-align: left;
            vertical-align: top;
            box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        }}
        
        .kpi-label {{
            font-family: 'Inter', sans-serif;
            font-size: 7.5pt;
            color: #7f8c8d;
            text-transform: uppercase;
            font-weight: 700;
            margin-bottom: 4px;
            display: block;
        }}
        
        .kpi-value {{
            font-family: 'Poppins', sans-serif;
            font-size: 14pt;
            font-weight: 700;
            color: #005151;
            margin: 0;
        }}
        
        .kpi-sub {{
            font-family: 'Inter', sans-serif;
            font-size: 7.5pt;
            font-weight: 600;
            margin-top: 3px;
            display: block;
        }}
        
        .kpi-separator {{
            border: 0;
            border-top: 1px solid #e2e8f0;
            margin: 10px 0 8px 0;
        }}
        
        .kpi-past-label {{
            font-family: 'Inter', sans-serif;
            font-size: 7pt;
            color: #a0aec0;
            text-transform: uppercase;
            font-weight: 700;
            margin-bottom: 2px;
            display: block;
        }}
        
        .kpi-past-value {{
            font-family: 'Poppins', sans-serif;
            font-size: 11pt;
            font-weight: 600;
            color: #718096;
            margin: 0;
        }}
        
        .negative {{ color: #c0392b; }}
        .positive {{ color: #159F92; }}

        /* Headers de Seção */
        .section-header {{
            font-family: 'Poppins', sans-serif;
            font-size: 10pt;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            color: #ffffff;
            background-color: #005151;
            padding: 5px 8px;
            margin-top: 18px;
            margin-bottom: 10px;
            font-weight: 700;
            border-radius: 2px;
            page-break-inside: avoid;
            page-break-after: avoid;
        }}
        
        /* Galeria de imagens adicionais */
        .image-gallery {{
            display: block;
            width: 100%;
            margin-top: 15px;
            margin-bottom: 15px;
            page-break-inside: avoid;
            text-align: center;
        }}
        
        .gallery-img {{
            display: inline-block;
            max-width: 100%;
            height: auto;
            margin: 5px 2px;
            border: 1px solid #e2e8f0;
            border-radius: 4px;
            padding: 4px;
            background: #ffffff;
            vertical-align: top;
            box-sizing: border-box;
        }}
        
        /* Expedia Grid Layout (2 rows, 1 col) */
        .expedia-gallery {{
            display: grid;
            grid-template-columns: 1fr;
            gap: 8px;
            width: 100%;
            margin-top: 15px;
        }}
        
        .expedia-gallery .grid-item-expedia {{
            width: 100%;
            text-align: center;
        }}
        
        .expedia-gallery .gallery-img {{
            max-width: 100%;
            max-height: 380px;
            width: 100%;
            margin: 0;
            border: 1px solid #e2e8f0;
            border-radius: 4px;
            padding: 4px;
            background: #ffffff;
            object-fit: contain;
            box-sizing: border-box;
        }}
        
        /* Booking Layout Table (Row 1 span 2, Row 2 split 70% and 30% - Highly compatible and robust) */
        .booking-table-layout {{
            width: 100% !important;
            height: 100% !important;
            border-collapse: collapse !important;
            border: none !important;
            margin-top: 15px !important;
            margin-bottom: 0 !important;
            table-layout: fixed !important;
        }}
        
        .booking-table-layout td {{
            border: none !important;
            background: none !important;
            padding: 0 !important;
            text-align: center !important;
            vertical-align: top !important;
        }}
        
        .gallery-img-booking-1 {{
            max-width: 100%;
            width: 100%;
            display: inline-block;
            border: 1px solid #e2e8f0;
            border-radius: 4px;
            padding: 1px;
            background: #ffffff;
            object-fit: contain;
            box-sizing: border-box;
        }}
        
        .gallery-img-booking-2,
        .gallery-img-booking-3 {{
            max-width: 100%;
            width: 100%;
            height: 380px;
            border: 1px solid #e2e8f0;
            border-radius: 4px;
            padding: 4px;
            background: #ffffff;
            object-fit: contain;
            box-sizing: border-box;
        }}
        
        /* Tabelas Financeiras */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
            font-size: 8.5pt;
            font-family: 'Inter', sans-serif;
            page-break-inside: avoid;
        }}
        
        th {{
            font-family: 'Poppins', sans-serif;
            font-weight: 700;
            text-transform: uppercase;
            font-size: 8pt;
            color: #ffffff;
            background-color: #159F92;
            padding: 6px 8px;
            text-align: right;
        }}
        
        th:first-child, td:first-child {{
            text-align: left;
        }}
        
        td {{
            font-family: 'Inter', sans-serif;
            padding: 5px 8px;
            border-bottom: 1px solid #e2e8f0;
            text-align: right;
            color: #2c3e50;
        }}

        tr:nth-child(even) td {{
            background-color: #f8f9fa;
        }}
        
        .total-row td {{
            font-weight: 700;
            background-color: #e8e8e8 !important;
            border-top: 1.5px solid #005151;
            border-bottom: 2px solid #005151;
        }}
    </style>
</head>
<body>

<div class="cover-page">
    <div class="cover-logo">
        <img src="{logo_url}" style="height: 180px; width: auto; display: block; margin: 0 auto;" />
    </div>
    
    <div class="cover-main">
        <h1 class="cover-hotel-name">{property_name}</h1>
        <h2 class="cover-doc-title">Lâmina de Performance Mensal</h2>
        <h3 class="cover-doc-reference">{reference}</h3>
    </div>
    
    <div class="cover-footer">
        <p class="cover-privacy-title">TERMO DE CONFIDENCIALIDADE E PRIVACIDADE</p>
        <p class="cover-privacy-text">
            Este documento contém informações comerciais confidenciais e proprietárias da <strong>Revinn Estratégias</strong> e do <strong>{property_name}</strong>. 
            O conteúdo aqui apresentado é destinado exclusivamente para uso interno e não deve ser reproduzido, compartilhado ou distribuído a terceiros sem a autorização prévia por escrito.
        </p>
        <p class="cover-date">São Paulo, 2026</p>
    </div>
</div>

<div class="container">

    <div class="table-row header-container">
        <div class="table-cell" style="width: 50%; vertical-align: middle;">
            <img src="{logo_url}" style="height: 100px; width: auto;" />
        </div>
        <div class="table-cell doc-info" style="width: 50%; vertical-align: middle;">
            <h1>{property_name}</h1>
            <p>Lâmina de Performance Mensal &bull; {reference}</p>
        </div>
    </div>

    {kpi_cards_html}

    <div class="section-header">Desempenho por Categoria de Acomodação (Check-outs)</div>
    <table>
        <thead>
            <tr>
                <th>Categoria de Quarto</th>
                <th>ADR Realizado</th>
                <th>LOS Média (Noites)</th>
                <th>Métricas de Alocação Absoluta</th>
            </tr>
        </thead>
        <tbody>
            {tabela_categorias_html}
        </tbody>
    </table>

    <div class="section-header">Atribuição de Performance por Canais de Emissão (Check-outs)</div>
    <table>
        <thead>
            <tr>
                <th>Canal de Distribuição</th>
                <th>Receita Bruta</th>
                <th>Reservas</th>
                <th>LOS</th>
                <th>ADR</th>
                <th>ADS (Janela)</th>
                <th>Share (%)</th>
            </tr>
        </thead>
        <tbody>
            {tabela_emissoes_passadas_html}
        </tbody>
    </table>

    <div class="section-header">Carteira de Captação Futura por Canais de Emissão (Pacing / OTB)</div>
    <table>
        <thead>
            <tr>
                <th>Canal Temático</th>
                <th>Volume Captado</th>
                <th>Reservas</th>
                <th>ADR Médio</th>
                <th>ADS (Antecedência)</th>
                <th>Share OTB (%)</th>
            </tr>
        </thead>
        <tbody>
            {tabela_emissoes_futuras_html}
        </tbody>
    </table>

    {additional_sections_html}

</div>

</body>
</html>"""

        # 5. Escrever PDF
        temp_html_path = output_path + ".temp.html"
        with open(temp_html_path, "w", encoding="utf-8") as f:
            f.write(html_template)
            
        try:
            HTML(temp_html_path).write_pdf(output_path)
        finally:
            if os.path.exists(temp_html_path):
                os.remove(temp_html_path)
